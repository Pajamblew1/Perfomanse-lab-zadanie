from docx import Document
import json
import re
import argparse
parse = argparse.ArgumentParser()
parse.add_argument("values")
parse.add_argument("tests")
parse.add_argument("report")
args = parse.parse_args()
diction = {}
doc1 = Document(args.values)
doc2 = Document(args.tests)
text1 = [paragraph.text for paragraph in doc1.paragraphs]
text2 = [paragraph.text for paragraph in doc2.paragraphs]
key = None
value1 = None
for paragraph in text1:
    if re.search(r'"id":',paragraph):
        key = paragraph.strip()
        
    if re.search(r'"value":',paragraph):
        value1 = paragraph.strip()
        diction[key]=value1
for i, paragraph in enumerate(text2):
    if re.search(r'"id":', paragraph):
        line = paragraph.strip()
        if line in diction:
            index = len(paragraph)-len(line)
            otstup = ' '*index
            text2[i+2]=otstup+diction[line]
with open(args.report, "w") as f:
    json.dump(text2, f, indent=2)

with open(args.report) as f:
    print(f.read())
